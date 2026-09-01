import os
import subprocess
import sys
import tempfile


def read_file(filepath: str) -> str:
    """
    Read an uploaded file.

    PDF files use the local PDF extraction pipeline.
    Text-based files are read directly.
    """

    if not filepath:
        return "Error: No filepath provided."

    if not os.path.exists(filepath):
        return f"File not found: {filepath}"

    try:
        extension = os.path.splitext(filepath)[1].lower()

        if extension == ".pdf":
            from vision.ocr import extract_text_from_document

            text = extract_text_from_document(filepath)

            if not text or not text.strip():
                return "Error: Could not extract any text from the PDF."

            return text.strip()

        with open(
            filepath,
            "r",
            encoding="utf-8",
            errors="ignore",
        ) as f:
            content = f.read()

        if not content.strip():
            return "The file is empty."

        return content

    except Exception as e:
        return f"Error reading file '{filepath}': {e}"


def search_kb(query: str) -> str:
    """
    Search the organization's local knowledge base using FAISS.
    """

    if not query or not query.strip():
        return "Error: Empty knowledge-base query."

    try:
        from rag.vector_store import search

        results = search(
            query,
            top_k=5,
        )

        if not results:
            return (
                "No relevant information was found in the "
                "organization knowledge base."
            )

        formatted_results = []

        for result in results:
            source = result.get(
                "source",
                "Unknown document",
            )

            text = result.get(
                "text",
                "",
            )

            score = result.get(
                "score",
                0,
            )

            formatted_results.append(
                f"Source: {source}\n"
                f"Relevance: {score:.3f}\n"
                f"Content:\n{text}"
            )

        return "\n\n---\n\n".join(
            formatted_results
        )

    except Exception as e:
        return (
            f"Error searching local knowledge base: {e}"
        )


def generate_docx(
    content: str,
    filename: str,
) -> str:
    """
    Generate a Word document using python-docx.
    """

    try:
        from docx import Document

    except ImportError:
        return (
            "Failed to create document: "
            "python-docx is not installed."
        )

    try:
        if not filename:
            filename = "Generated_Document.docx"

        if not filename.lower().endswith(".docx"):
            filename += ".docx"

        directory = os.path.dirname(filename)

        if directory:
            os.makedirs(
                directory,
                exist_ok=True,
            )

        document = Document()

        if content and content.strip():

            for paragraph in content.split("\n\n"):

                paragraph = paragraph.strip()

                if paragraph:
                    document.add_paragraph(
                        paragraph
                    )

        else:
            document.add_paragraph(
                "No document content was generated."
            )

        document.save(filename)

        return (
            f"Created {filename} successfully."
        )

    except Exception as e:
        return (
            f"Failed to create {filename}: {e}"
        )


def run_code_in_sandbox(
    code: str,
) -> str:
    """
    Execute Python code.

    This is currently a local fallback.
    The Docker sandbox can be connected later.
    """

    if not code or not code.strip():
        return (
            "Code execution failed: "
            "No code provided."
        )

    script_path = None

    try:

        with tempfile.NamedTemporaryFile(
            mode="w",
            suffix=".py",
            delete=False,
            encoding="utf-8",
        ) as f:

            f.write(code)

            script_path = f.name

        result = subprocess.run(
            [
                sys.executable,
                script_path,
            ],
            capture_output=True,
            text=True,
            timeout=10,
        )

        if result.returncode == 0:

            output = result.stdout.strip()

            if not output:
                output = "(No output)"

            return (
                "Code executed successfully.\n"
                f"Output:\n{output}"
            )

        error = result.stderr.strip()

        if not error:
            error = (
                "Unknown execution error."
            )

        return (
            f"Code execution failed "
            f"(exit {result.returncode}).\n"
            f"Error:\n{error}"
        )

    except subprocess.TimeoutExpired:

        return (
            "Code execution timed out "
            "after 10 seconds."
        )

    except Exception as e:

        return (
            f"Sandbox execution error: {e}"
        )

    finally:

        if (
            script_path
            and os.path.exists(script_path)
        ):

            try:
                os.unlink(script_path)

            except Exception:
                pass